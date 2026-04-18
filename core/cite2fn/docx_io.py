"""Document I/O utilities.

Loading, saving, and section removal for .docx documents.
"""

from __future__ import annotations

import shutil
from pathlib import Path
from docx import Document


def load_document(path: str) -> Document:
    """Load a .docx document."""
    return Document(path)


def save_document(doc: Document, output_path: str) -> None:
    """Save a document to a new file."""
    doc.save(output_path)


def copy_document(input_path: str, output_path: str) -> None:
    """Copy a .docx file to work on a copy (non-destructive)."""
    shutil.copy2(input_path, output_path)


def remove_references_section(doc: Document) -> int:
    """Remove the References/Bibliography section from the document.

    Returns the number of paragraphs removed.
    """
    from cite2fn.references import find_references_section

    start, end = find_references_section(doc)
    if start is None:
        return 0

    # Remove paragraphs from end to start to avoid index shifting
    body = doc.element.body
    paragraphs_to_remove = []

    for i in range(start, min(end, len(doc.paragraphs))):
        paragraphs_to_remove.append(doc.paragraphs[i]._element)

    for p_elem in paragraphs_to_remove:
        body.remove(p_elem)

    return len(paragraphs_to_remove)
