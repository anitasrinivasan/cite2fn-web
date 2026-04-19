"""Inline text cleanup after footnote insertion.

Implements Rules 1-5 from the spec:
1. Author name is grammatical → keep name, drop year
2. Standalone parenthetical → remove entirely
3. Entire hyperlinked phrase is citation (no grammatical role) → remove
4. Hyperlink on non-citation text → remove hyperlink formatting, keep text
5. Remove hyperlink formatting after adding footnote
"""

from __future__ import annotations

import re
from lxml import etree
from docx import Document

from cite2fn.models import Citation

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def classify_cleanup_rule(citation: Citation, para_text: str) -> int:
    """Determine which cleanup rule applies to a citation.

    Returns the rule number (1-5).
    """
    display = citation.display_text

    if citation.type == "parenthetical":
        # Check if the author name appears before the parenthetical
        # e.g., "Bisbee et al. (2023) found" → Rule 1 (inline author + year in parens)
        # vs "(Spann et al., 2025)" → Rule 2 (standalone)

        # Find the citation in the text
        idx = para_text.find(display)
        if idx >= 0:
            before = para_text[:idx].rstrip()
            # If the parenthetical is at the start or preceded by a period/sentence boundary
            if not before or before[-1] in ".!?:;":
                return 2  # Standalone at start
            # If preceded by author name text that's part of the sentence
            return 2  # Default for parenthetical: remove entirely

    if citation.type == "inline_author_date":
        # Keep author name only if it's grammatically part of a sentence —
        # i.e. there's substantive text continuing after the citation. If it's
        # dangling at the end of a paragraph or followed only by punctuation/
        # whitespace (e.g. "Hämäläinen et al. (2023)" with nothing after),
        # the author isn't a subject/object of anything and should go with
        # the year.
        idx = para_text.find(display)
        if idx >= 0:
            after = para_text[idx + len(display):]
            if not re.search(r"\w", after):
                return 2
        return 1

    if citation.type in ("hyperlink_external", "hyperlink_internal"):
        idx = para_text.find(display)
        if idx >= 0:
            before_full = para_text[:idx]
            after_full = para_text[idx + len(display):]
            before = before_full.rstrip()
            after = after_full.lstrip()

            # Rule 2 (remove entirely) when the hyperlink IS itself a standalone
            # citation — either its display already has parens, or the surrounding
            # text wraps it in parens with no other content between the parens.
            if display.strip().startswith("("):
                return 2
            if before.endswith("(") and after.startswith(")"):
                return 2

            # Non-citation hyperlink (e.g. a project/product name): no year
            # signal at all.
            if not citation.year and "et al" not in display.lower():
                if not re.search(r"\d{4}", display):
                    return 4  # Keep text, remove hyperlink formatting

            # Everything else: author is a grammatical part of the sentence.
            # Keep the author name, drop the year. This replaces the earlier
            # "before ends in punctuation → rule 3" heuristic which was
            # over-aggressive: e.g. "results. Jackaria et al (2024) found…"
            # would delete the author even though it's the subject of "found".
            return 1

    if citation.type == "existing_footnote":
        return 0  # No cleanup needed for existing footnotes

    return 5  # Default: just remove hyperlink formatting


def apply_cleanup(
    doc: Document,
    citation: Citation,
    rule: int,
) -> None:
    """Apply the cleanup rule to the document for a given citation.

    Modifies the document XML in-place.
    """
    if rule == 0:
        return  # No cleanup needed

    if citation.type in ("parenthetical", "inline_author_date"):
        _cleanup_text_citation(doc, citation, rule)
    elif citation.type in ("hyperlink_external", "hyperlink_internal"):
        _cleanup_hyperlink_citation(doc, citation, rule)

    # Paragraph-wide whitespace tidy-up after removal: removing "(X 2025)"
    # from "text (X 2025). Next" leaves "text . Next" — collapse the stray
    # space-before-punctuation and any resulting double space.
    para = doc.paragraphs[citation.paragraph_index]
    _collapse_stray_whitespace(para._element)


def _collapse_stray_whitespace(para_elem: etree._Element) -> None:
    """Collapse patterns produced by removing a citation mid-sentence.

    - " ." / " ," / " ;" / " :" / " !" / " ?" / " )" / " ]"  → remove the space
    - "( " / "[ "                                             → remove the space
    - "  " (double space) → single space
    Operates on each `w:t` independently but also handles the "end of one run,
    start of the next" boundary by checking the last char of each text against
    the first char of the next text.
    """
    t_elems = para_elem.findall(f".//{{{W}}}t")

    # Pass 1: within-text patterns.
    for t in t_elems:
        if not t.text:
            continue
        new = t.text
        new = re.sub(r" +([.,;:!?\)\]])", r"\1", new)
        new = re.sub(r"([\(\[]) +", r"\1", new)
        new = re.sub(r"  +", " ", new)
        if new != t.text:
            t.text = new

    # Pass 2: cross-run boundaries. If one text ends with " " and the next
    # starts with punctuation, or one ends with "(" and next starts with " ",
    # trim.
    populated = [t for t in t_elems if t.text]
    for i in range(len(populated) - 1):
        cur = populated[i]
        nxt = populated[i + 1]
        if cur.text.endswith(" ") and nxt.text and nxt.text[0] in ".,;:!?)]":
            cur.text = cur.text.rstrip(" ")
        if cur.text.endswith("(") or cur.text.endswith("["):
            if nxt.text.startswith(" "):
                nxt.text = nxt.text.lstrip(" ")
        # double-space across runs
        if cur.text.endswith(" ") and nxt.text.startswith(" "):
            nxt.text = nxt.text.lstrip(" ")


def _cleanup_text_citation(doc: Document, citation: Citation, rule: int) -> None:
    """Clean up parenthetical or inline author-date citations in body text."""
    para = doc.paragraphs[citation.paragraph_index]
    para_elem = para._element

    if rule == 2:
        # Remove the entire parenthetical text
        _remove_text_from_paragraph(para_elem, citation.display_text)
    elif rule == 1:
        # Keep author name, remove year within the citation's display_text span.
        # Must scope to the display_text span: multiple citations in the same
        # paragraph can share year patterns (e.g. both Kaiser's "(2025)" and
        # Wang inline's "(2025)" exist simultaneously during cleanup), and a
        # paragraph-wide search would yank the wrong one.
        if citation.year:
            _remove_year_within_display(para_elem, citation.display_text, citation.year)


def _cleanup_hyperlink_citation(doc: Document, citation: Citation, rule: int) -> None:
    """Clean up hyperlinked citations by removing hyperlink formatting."""
    para = doc.paragraphs[citation.paragraph_index]
    para_elem = para._element

    # Find the hyperlink element(s) for this citation
    hyperlinks = para_elem.findall(f".//{{{W}}}hyperlink")

    for hl in hyperlinks:
        texts = [t.text for t in hl.findall(f".//{{{W}}}t") if t.text]
        hl_text = "".join(texts)

        if not _text_matches(hl_text, citation.display_text):
            continue

        if rule in (2, 3):
            # Remove entirely — remove the hyperlink and all its content.
            # Rule 2 = standalone parenthetical style citation (wrapped in
            # parens in the body). Rule 3 = hyperlinked phrase that IS the
            # citation. Both need the same sweep: drop the hyperlink, any
            # orphan brackets that were inside it, and any matching
            # wrapping parens in the surrounding runs.
            parent = hl.getparent()
            hl_idx = list(parent).index(hl)
            parent.remove(hl)
            _remove_orphan_brackets(parent, hl_idx, hl_text)
            _remove_wrapping_brackets(parent, hl_idx)
        elif rule in (1, 4, 5):
            # Keep text, remove hyperlink wrapper
            parent = hl.getparent()
            hl_idx = list(parent).index(hl)
            # Capture the hyperlink's text runs BEFORE unwrap so we can
            # target year-removal at just these runs. Scanning the whole
            # paragraph would accidentally eat another nearby citation's year
            # (e.g. Bhatnagar's cleanup would otherwise strip Pataranutaporn's
            # "(2025)" because the paragraph-level search matches the earliest
            # occurrence, not the one belonging to this citation).
            hl_runs = list(hl)
            _unwrap_hyperlink(hl)

            if rule == 1:
                # Scoped year removal — search within the citation's
                # display_text span in the paragraph. Covers the case where
                # detection captured a partial display (e.g.
                # "Vaishampayan et al. (2025" with ")" in the next run)
                # because the paragraph-level combined text still has the
                # full " (YYYY)".
                if citation.year:
                    _remove_year_within_display(para_elem, citation.display_text, citation.year)
                # Handle the separate case where detection's display_text
                # omitted the year entirely (e.g. "Iso et al.") and the
                # "(YYYY)" fragment lives in a run AFTER the unwrapped
                # hyperlink position.
                _strip_trailing_year_parenthetical(parent, hl_idx + len(hl_runs))


def _unwrap_hyperlink(hyperlink: etree._Element) -> None:
    """Remove hyperlink wrapper, keeping the child runs in place."""
    parent = hyperlink.getparent()
    idx = list(parent).index(hyperlink)

    # Move all children out of the hyperlink
    children = list(hyperlink)
    for i, child in enumerate(children):
        hyperlink.remove(child)
        parent.insert(idx + i, child)

        # Remove hyperlink-style formatting from the runs
        rpr = child.find(f"{{{W}}}rPr")
        if rpr is not None:
            # Remove color
            for color in rpr.findall(f"{{{W}}}color"):
                rpr.remove(color)
            # Remove underline
            for u in rpr.findall(f"{{{W}}}u"):
                rpr.remove(u)
            # Remove hyperlink style
            for rstyle in rpr.findall(f"{{{W}}}rStyle"):
                if rstyle.get(f"{{{W}}}val") in ("Hyperlink", "InternetLink"):
                    rpr.remove(rstyle)

    parent.remove(hyperlink)


_BRACKET_PAIRS = {"(": ")", "[": "]", "{": "}", "“": "”", "\"": "\""}


def _strip_trailing_year_parenthetical(parent: etree._Element, start_idx: int) -> None:
    """After a rule-1 hyperlink unwrap, scan forward from `start_idx` and
    strip a `(YYYY)` or ` (YYYY)` fragment if it's the first non-trivial
    content encountered.

    Used when detection's display_text didn't capture the year (because the
    parenthetical year was outside the hyperlink's runs). Only touches the
    first non-whitespace text node — we never look past the first real word.
    """
    pattern = re.compile(r"^(\s*)\(?\s*(\d{4})\s*\)?([,.;:]?)")
    siblings = list(parent)
    for i in range(start_idx, len(siblings)):
        sib = siblings[i]
        if sib.tag != f"{{{W}}}r":
            continue
        t_elems = sib.findall(f"{{{W}}}t")
        if not t_elems:
            continue
        text = "".join(t.text or "" for t in t_elems)
        if not text.strip():
            continue  # skip empty/whitespace-only runs

        # Must start with ` (YYYY)` or `(YYYY)` (possibly preceded by ws).
        lead_ws_match = re.match(r"(\s*)", text)
        lead_ws = lead_ws_match.group(1) if lead_ws_match else ""
        remainder = text[len(lead_ws):]
        m = re.match(r"\(\s*(\d{4})\s*\)", remainder)
        if not m:
            return  # not a year-parenthetical; stop scanning

        # Strip the matched "(YYYY)" from the text, keeping any leading ws.
        new_text = lead_ws + remainder[m.end():]
        # Spread across the original t_elems: simplest — put all in first, blank rest
        t_elems[0].text = new_text
        t_elems[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        for t in t_elems[1:]:
            t.text = ""
        return


def _remove_wrapping_brackets(
    parent: etree._Element, removed_idx: int
) -> None:
    """After removing a hyperlink (or any citation element), look for a matched
    bracket pair that WRAPS the removal site — `(` in the preceding text and
    `)` in the following text — and remove both. Ignores whitespace between
    the removal site and the brackets.

    Handles `()`, `[]`, `{}`. Does not touch quote marks (too ambiguous).
    """
    siblings = list(parent)

    # Find the nearest non-empty text character before the removal position.
    prev_ctx = _last_char_before(siblings, removed_idx)
    next_ctx = _first_char_after(siblings, removed_idx)

    if prev_ctx is None or next_ctx is None:
        return

    prev_t, prev_char = prev_ctx
    next_t, next_char = next_ctx

    wrappers = {"(": ")", "[": "]", "{": "}"}
    if prev_char in wrappers and wrappers[prev_char] == next_char:
        # Drop the last `(` from prev_t and the first `)` from next_t,
        # leaving surrounding whitespace alone (collapsed later by
        # _collapse_stray_whitespace).
        pidx = prev_t.text.rfind(prev_char)
        if pidx >= 0:
            prev_t.text = prev_t.text[:pidx] + prev_t.text[pidx + 1:]
        nidx = next_t.text.find(next_char)
        if nidx >= 0:
            next_t.text = next_t.text[:nidx] + next_t.text[nidx + 1:]


def _last_char_before(siblings: list, idx: int):
    """Return (w:t elem, last visible char) for the run before `idx`, or None.
    Skips fully-whitespace text nodes."""
    for i in range(idx - 1, -1, -1):
        sib = siblings[i]
        if sib.tag != f"{{{W}}}r":
            continue
        t_elems = sib.findall(f"{{{W}}}t")
        for t in reversed(t_elems):
            if t.text and t.text.strip():
                # Trim trailing whitespace to find the last non-space char,
                # but keep returning the t element itself.
                stripped = t.text.rstrip()
                return (t, stripped[-1])
    return None


def _first_char_after(siblings: list, idx: int):
    """Return (w:t elem, first visible char) for the run at or after `idx`, or None."""
    for i in range(idx, len(siblings)):
        sib = siblings[i]
        if sib.tag != f"{{{W}}}r":
            continue
        t_elems = sib.findall(f"{{{W}}}t")
        for t in t_elems:
            if t.text and t.text.strip():
                stripped = t.text.lstrip()
                return (t, stripped[0])
    return None


def _remove_orphan_brackets(
    parent: etree._Element, removed_idx: int, removed_text: str
) -> None:
    """After a rule-3 hyperlink removal, strip a lone orphan bracket that was
    paired with one inside the hyperlink.

    Counts opens vs closes in `removed_text`:
    - More opens than closes → removed text contained an unpaired `(`. The
      matching `)` is somewhere to the right of the removal point; scan the
      immediately-following run and strip a leading `)` if present.
    - More closes than opens → unpaired `)`. Strip a trailing `(` from the
      preceding run.
    Handles `()`, `[]`, `{}`, straight and curly quotes.
    """
    for opener, closer in _BRACKET_PAIRS.items():
        open_count = removed_text.count(opener)
        close_count = removed_text.count(closer)
        if opener == closer:  # quote marks — can't tell open from close
            if open_count % 2 == 1:
                # Odd count → unpaired. Try stripping from the next run first,
                # then from the previous.
                if not _strip_leading_char(parent, removed_idx, closer):
                    _strip_trailing_char(parent, removed_idx - 1, opener)
            continue
        if open_count > close_count:
            _strip_leading_char(parent, removed_idx, closer)
        elif close_count > open_count:
            _strip_trailing_char(parent, removed_idx - 1, opener)


def _strip_leading_char(parent: etree._Element, idx: int, char: str) -> bool:
    """Remove `char` from the start of the first text content at or after idx.
    Returns True if a character was actually stripped."""
    siblings = list(parent)
    for i in range(idx, len(siblings)):
        sib = siblings[i]
        if sib.tag != f"{{{W}}}r":
            continue
        for t in sib.findall(f"{{{W}}}t"):
            if t.text:
                if t.text.startswith(char):
                    t.text = t.text[1:]
                    return True
                return False
    return False


def _strip_trailing_char(parent: etree._Element, idx: int, char: str) -> bool:
    """Remove `char` from the end of the last text content at or before idx.
    Returns True if a character was actually stripped."""
    siblings = list(parent)
    for i in range(min(idx, len(siblings) - 1), -1, -1):
        sib = siblings[i]
        if sib.tag != f"{{{W}}}r":
            continue
        t_elems = sib.findall(f"{{{W}}}t")
        if not t_elems:
            continue
        last = t_elems[-1]
        if last.text:
            if last.text.endswith(char):
                last.text = last.text[:-1]
                return True
            return False
    return False


def _text_matches(hl_text: str, citation_text: str) -> bool:
    """Check if hyperlink text matches a citation's display text (fuzzy)."""
    # Normalize whitespace and strip
    hl_norm = " ".join(hl_text.split()).strip("()[].,; ")
    cite_norm = " ".join(citation_text.split()).strip("()[].,; ")
    # Check containment in either direction
    return hl_norm in cite_norm or cite_norm in hl_norm or hl_norm == cite_norm


def _remove_text_from_paragraph(para_elem: etree._Element, text_to_remove: str) -> bool:
    """Remove a specific text string from paragraph runs. Returns True on success."""
    for run in para_elem.findall(f".//{{{W}}}r"):
        for t_elem in run.findall(f"{{{W}}}t"):
            if t_elem.text and text_to_remove in t_elem.text:
                t_elem.text = t_elem.text.replace(text_to_remove, "", 1)
                # Clean up double spaces
                t_elem.text = re.sub(r"  +", " ", t_elem.text)
                return True

    # If not found in a single run, try across consecutive runs
    # (the text may span multiple runs)
    runs = para_elem.findall(f".//{{{W}}}r")
    full_text = ""
    run_map: list[tuple[etree._Element, int, int]] = []  # (t_elem, start, end)
    for run in runs:
        for t_elem in run.findall(f"{{{W}}}t"):
            if t_elem.text:
                start = len(full_text)
                full_text += t_elem.text
                run_map.append((t_elem, start, len(full_text)))

    idx = full_text.find(text_to_remove)
    if idx < 0:
        return False
    remove_end = idx + len(text_to_remove)
    for t_elem, start, end in run_map:
        if end <= idx or start >= remove_end:
            continue
        # This run overlaps with the text to remove
        text = t_elem.text
        local_start = max(0, idx - start)
        local_end = min(len(text), remove_end - start)
        t_elem.text = text[:local_start] + text[local_end:]
    return True


def _remove_year_within_display(
    para_elem: etree._Element, display_text: str, year: str
) -> bool:
    """Find `display_text` in the paragraph and remove the year fragment
    (` (YYYY)`, `(YYYY)`, `, YYYY`, or ` YYYY`) only within that span.

    Prevents cross-citation collision: Wang inline's cleanup must only touch
    Wang inline's own "(2025)", not Kaiser's "(2025)" that happens to appear
    earlier in the same paragraph.
    """
    patterns = [f" ({year})", f"({year})", f", {year}", f" {year}"]

    combined = ""
    run_map: list[tuple[etree._Element, int, int]] = []
    for r in para_elem.findall(f".//{{{W}}}r"):
        for t_elem in r.findall(f"{{{W}}}t"):
            if t_elem.text:
                start = len(combined)
                combined += t_elem.text
                run_map.append((t_elem, start, len(combined)))

    disp_idx = combined.find(display_text)
    if disp_idx < 0:
        return False
    disp_end = disp_idx + len(display_text)

    for pat in patterns:
        pat_idx = combined.find(pat, disp_idx)
        # Pattern must START within the display span (or immediately after
        # when detection truncated the display — e.g. "Vaishampayan et al.
        # (2025" with no closing paren captured). Allow the match to extend a
        # few chars past display_end to catch the ")".
        if pat_idx < 0 or pat_idx > disp_end:
            continue
        pat_end = pat_idx + len(pat)
        for t_elem, s, e in run_map:
            if e <= pat_idx or s >= pat_end:
                continue
            local_s = max(0, pat_idx - s)
            local_e = min(len(t_elem.text or ""), pat_end - s)
            t_elem.text = (t_elem.text or "")[:local_s] + (t_elem.text or "")[local_e:]
        return True
    return False


def _remove_year_from_specific_runs(runs: list, year: str) -> bool:
    """Remove a ` (YYYY)`, `(YYYY)`, `, YYYY`, or ` YYYY` pattern found only
    within the given list of w:r elements. Does NOT scan siblings. Used by
    hyperlink rule-1 cleanup so one citation's year-removal doesn't eat
    another citation's year elsewhere in the paragraph.

    Returns True if a pattern was removed.
    """
    patterns = [f" ({year})", f"({year})", f", {year}", f" {year}"]

    # Concatenate text across all given runs + track which t_elem owns which
    # offset range.
    combined = ""
    run_map: list[tuple[etree._Element, int, int]] = []
    for r in runs:
        for t_elem in r.findall(f"{{{W}}}t"):
            if t_elem.text:
                start = len(combined)
                combined += t_elem.text
                run_map.append((t_elem, start, len(combined)))

    for pat in patterns:
        idx = combined.find(pat)
        if idx < 0:
            continue
        end = idx + len(pat)
        for t_elem, s, e in run_map:
            if e <= idx or s >= end:
                continue
            local_s = max(0, idx - s)
            local_e = min(len(t_elem.text or ""), end - s)
            t_elem.text = (t_elem.text or "")[:local_s] + (t_elem.text or "")[local_e:]
        return True
    return False


def _remove_year_from_runs(
    para_elem: etree._Element, year: str, context: str
) -> None:
    """Remove year text (and surrounding parens) from runs near a citation.

    Tries patterns in decreasing specificity and stops on the first successful
    removal. Without the bool return from _remove_text_from_paragraph this
    function would fall through all patterns when the year's run layout doesn't
    match the first guess.
    """
    patterns = [f" ({year})", f"({year})", f", {year}", f" {year}"]
    for pat in patterns:
        if _remove_text_from_paragraph(para_elem, pat):
            return
