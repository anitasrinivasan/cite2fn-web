"""Citation detection across all 5 patterns.

Scans a .docx document for:
1. Hyperlinked text with external URL
2. Hyperlinked text with internal anchor (references section)
3. Plain parenthetical citations (no hyperlink)
4. Inline author-date citations (no hyperlink)
5. Existing footnotes containing URLs or non-Bluebook text
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from docx import Document
from lxml import etree

from cite2fn.models import Citation

# Namespaces used in Word XML
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NSMAP = {"w": W, "r": R}

# Google Docs URL pattern — these are internal references, not fetchable URLs
GOOGLE_DOCS_RE = re.compile(r"docs\.google\.com/document/")

# Author name pattern: handles "Smith", "DellaVigna", "Bar-Gill", "O'Brien"
_AUTHOR = r"[A-ZÀ-Ž][a-zà-ž'\-]*(?:[A-Z][a-zà-ž'\-]+)*"

# Regex for parenthetical citations: (Author et al., 2023) and variants
# Handles: (Smith 2020), (Smith, 2020), (Smith et al. 2020), (Smith et al., 2020),
# (Smith & Jones, 2020), (Smith, Jones, and Doe, 2020), (Smith and Jones 2020)
PARENTHETICAL_RE = re.compile(
    r"\("
    r"("
    + _AUTHOR +  # First author last name
    r"(?:\s+et\s+al\.?)?"  # Optional "et al."
    r"(?:,\s*" + _AUTHOR + r")*"  # Comma-separated additional names
    r"(?:[,\s]+(?:and|&)\s+" + _AUTHOR + r")?"  # Final "and/& Author"
    r",?\s*\d{4}"  # Year
    r")"
    r"\)",
    re.UNICODE,
)

# Regex for inline author-date: Author et al. (2023), Author (2023), Author and Author (2023)
INLINE_AUTHOR_DATE_RE = re.compile(
    r"(" + _AUTHOR +
    r"(?:\s+et\s+al\.?)?"  # "et al." / "et al"
    r"(?:\s+(?:and|&)\s+" + _AUTHOR + r")?"  # "and Author" / "& Author"
    r")"
    r"\s*\((\d{4})\)",
    re.UNICODE,
)

# Regex for inline author-date WITHOUT parens: "Hullman 2025 describes"
INLINE_NO_PARENS_RE = re.compile(
    r"(" + _AUTHOR +
    r"(?:\s+et\s+al\.?)?"
    r"(?:\s+(?:and|&)\s+" + _AUTHOR + r")?"
    r")"
    r"\s+(\d{4})"
    r"(?=[\s,;.\)]|$)",  # Year must be followed by space, punct, or end
    re.UNICODE,
)

# Words that look like author names but aren't (common false positives)
_FALSE_POSITIVE_AUTHORS = {
    "In", "The", "For", "From", "With", "About", "Since", "During",
    "After", "Before", "Between", "Under", "Over", "Part", "Section",
    "Chapter", "Table", "Figure", "Vol", "No", "January", "February",
    "March", "April", "May", "June", "July", "August", "September",
    "October", "November", "December",
}

# URL pattern for detecting URLs in footnote text
URL_RE = re.compile(r"https?://[^\s<>\"]+")

# Common Bluebook signals
SIGNAL_RE = re.compile(
    r"\b(See(?:\s+also)?|Cf\.|Compare|But\s+see|See\s+generally|E\.g\.,?)\s*$",
    re.IGNORECASE,
)


def _get_paragraph_text(para_elem: etree._Element) -> str:
    """Extract full text from a paragraph XML element, including hyperlink text."""
    texts = []
    for t in para_elem.findall(f".//{{{W}}}t"):
        if t.text:
            texts.append(t.text)
    return "".join(texts)


def _get_surrounding_sentence(full_text: str, match_start: int, match_end: int) -> str:
    """Extract the sentence surrounding a match position."""
    # Find sentence boundaries (period, question mark, exclamation)
    sent_start = max(0, full_text.rfind(". ", 0, match_start) + 2)
    sent_end = full_text.find(". ", match_end)
    if sent_end == -1:
        sent_end = len(full_text)
    else:
        sent_end += 1
    return full_text[sent_start:sent_end].strip()


@dataclass
class _RawHyperlink:
    """Intermediate representation of a hyperlink found in the XML."""
    para_index: int
    element: etree._Element
    display_text: str
    url: str | None
    anchor: str | None
    position_in_para: int  # child index within paragraph


def _find_references_section_start(doc: Document) -> int | None:
    """Find the paragraph index where the References/Bibliography section starts.

    Returns None if no such section exists.
    """
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else ""
        # Check heading styles or standalone "References" / "Bibliography" text
        is_heading = "heading" in style.lower() or "Heading" in style
        if text.lower() in ("references", "bibliography", "works cited"):
            if is_heading or len(text) < 20:
                return i
    return None


def detect_citations(doc_path: str) -> list[Citation]:
    """Detect all citations in a .docx document.

    Returns a list of Citation objects covering all 5 patterns.
    """
    doc = Document(doc_path)
    citations: list[Citation] = []
    cite_id = 0

    def next_id() -> str:
        nonlocal cite_id
        cite_id += 1
        return f"cite-{cite_id:03d}"

    # Find where the References section starts (exclude from body detection)
    refs_start = _find_references_section_start(doc)

    # --- Pattern 1 & 2: Hyperlinks ---
    citations.extend(_detect_hyperlinks(doc, next_id, refs_start))

    # --- Pattern 3 & 4: Parenthetical and inline author-date ---
    # Collect hyperlink text spans to avoid double-detection
    hyperlink_spans = _get_hyperlink_text_spans(doc)
    citations.extend(_detect_parenthetical(doc, next_id, hyperlink_spans, refs_start))
    citations.extend(_detect_inline_author_date(doc, next_id, hyperlink_spans, refs_start))

    # --- Pattern 5: Existing footnotes ---
    citations.extend(_detect_existing_footnotes(doc, next_id))

    return citations


def _detect_hyperlinks(doc: Document, next_id, refs_start: int | None = None) -> list[Citation]:
    """Detect hyperlinked citations (Patterns 1 and 2)."""
    citations = []

    for para_idx, para in enumerate(doc.paragraphs):
        if refs_start is not None and para_idx >= refs_start:
            break
        para_elem = para._element
        para_text = para.text

        # Collect all hyperlinks in this paragraph
        raw_links: list[_RawHyperlink] = []
        for child_idx, child in enumerate(para_elem):
            tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ""
            if tag == "hyperlink":
                anchor = child.get(f"{{{W}}}anchor")
                rid = child.get(f"{{{R}}}id")
                url = None
                if rid:
                    try:
                        url = doc.part.rels[rid].target_ref
                    except KeyError:
                        pass

                texts = [t.text for t in child.findall(f".//{{{W}}}t") if t.text]
                display = "".join(texts)

                raw_links.append(_RawHyperlink(
                    para_index=para_idx,
                    element=child,
                    display_text=display,
                    url=url,
                    anchor=anchor,
                    position_in_para=child_idx,
                ))

        # Merge adjacent hyperlinks sharing the same target
        merged = _merge_adjacent_hyperlinks(raw_links)

        for group in merged:
            display = "".join(rl.display_text for rl in group)
            url = group[0].url
            anchor = group[0].anchor

            # Classify: external vs internal
            is_internal = anchor is not None
            if url and GOOGLE_DOCS_RE.search(url):
                is_internal = True

            cite_type = "hyperlink_internal" if is_internal else "hyperlink_external"

            # Parse author and year from display text
            author, year = _parse_author_year(display)

            citations.append(Citation(
                id=next_id(),
                type=cite_type,
                display_text=display.strip(),
                paragraph_index=para_idx,
                surrounding_sentence=_get_surrounding_sentence(
                    para_text, 0, len(para_text)
                ),
                url=url if not is_internal else None,
                internal_anchor=anchor,
                author_name=author,
                year=year,
                hyperlink_index=group[0].position_in_para,
            ))

    return citations


def _merge_adjacent_hyperlinks(
    raw_links: list[_RawHyperlink],
) -> list[list[_RawHyperlink]]:
    """Merge consecutive hyperlinks that share the same target (anchor or URL).

    Paper 1 has split hyperlinks like:
      <w:hyperlink anchor="_hqe1036ucurq">(Rossi et al.,</w:hyperlink>
      <w:hyperlink anchor="_hqe1036ucurq">1996).</w:hyperlink>
    These need to be treated as one citation.
    """
    if not raw_links:
        return []

    groups: list[list[_RawHyperlink]] = [[raw_links[0]]]

    for rl in raw_links[1:]:
        prev = groups[-1][-1]
        same_target = False
        if rl.anchor and prev.anchor and rl.anchor == prev.anchor:
            same_target = True
        elif rl.url and prev.url and rl.url == prev.url:
            same_target = True

        # Check adjacency: position should differ by a small amount
        # (there may be intervening w:r elements with spaces)
        adjacent = (rl.position_in_para - prev.position_in_para) <= 3

        if same_target and adjacent:
            groups[-1].append(rl)
        else:
            groups.append([rl])

    return groups


def _parse_author_year(text: str) -> tuple[str | None, str | None]:
    """Extract author name and year from citation display text."""
    text = text.strip().strip("()[].,;")

    _auth = r"[A-ZÀ-Ž][a-zà-ž'\-]*(?:[A-Z][a-zà-ž'\-]+)*"

    # Try "Author et al. (YYYY)" or "Author et al., YYYY" or "Author & Author, YYYY"
    m = re.search(
        r"(" + _auth + r"(?:\s+(?:et\s+al\.?|and|&),?\s+" + _auth + r")?)"
        r".*?(\d{4})",
        text,
        re.UNICODE,
    )
    if m:
        return m.group(1).strip(), m.group(2)

    # Try just author name (no year)
    m = re.match(r"(" + _auth + r"(?:\s+(?:et\s+al\.?))?)", text, re.UNICODE)
    if m:
        return m.group(1).strip(), None

    return None, None


def _get_hyperlink_text_spans(doc: Document) -> dict[int, list[tuple[int, int]]]:
    """Get character spans of hyperlinked text in each paragraph.

    Returns {paragraph_index: [(start, end), ...]} so we can avoid
    double-detecting hyperlinked text as parenthetical/inline citations.
    """
    spans: dict[int, list[tuple[int, int]]] = {}

    for para_idx, para in enumerate(doc.paragraphs):
        para_elem = para._element
        pos = 0
        para_spans = []

        for child in para_elem:
            tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ""

            if tag == "hyperlink":
                texts = [t.text for t in child.findall(f".//{{{W}}}t") if t.text]
                hl_text = "".join(texts)
                para_spans.append((pos, pos + len(hl_text)))
                pos += len(hl_text)
            elif tag == "r":
                texts = [t.text for t in child.findall(f".//{{{W}}}t") if t.text]
                pos += sum(len(t) for t in texts)

        if para_spans:
            spans[para_idx] = para_spans

    return spans


def _overlaps_hyperlink(
    para_idx: int,
    match_start: int,
    match_end: int,
    hyperlink_spans: dict[int, list[tuple[int, int]]],
) -> bool:
    """Check if a text match overlaps with any hyperlink span."""
    if para_idx not in hyperlink_spans:
        return False
    for hl_start, hl_end in hyperlink_spans[para_idx]:
        if match_start < hl_end and match_end > hl_start:
            return True
    return False


def _detect_parenthetical(
    doc: Document,
    next_id,
    hyperlink_spans: dict[int, list[tuple[int, int]]],
    refs_start: int | None = None,
) -> list[Citation]:
    """Detect parenthetical citations like (Smith et al., 2023)."""
    citations = []

    for para_idx, para in enumerate(doc.paragraphs):
        if refs_start is not None and para_idx >= refs_start:
            break
        text = para.text
        if not text:
            continue

        for m in PARENTHETICAL_RE.finditer(text):
            if _overlaps_hyperlink(para_idx, m.start(), m.end(), hyperlink_spans):
                continue

            inner = m.group(1)
            author, year = _parse_author_year(inner)

            citations.append(Citation(
                id=next_id(),
                type="parenthetical",
                display_text=m.group(0),
                paragraph_index=para_idx,
                surrounding_sentence=_get_surrounding_sentence(
                    text, m.start(), m.end()
                ),
                author_name=author,
                year=year,
            ))

    return citations


def _detect_inline_author_date(
    doc: Document,
    next_id,
    hyperlink_spans: dict[int, list[tuple[int, int]]],
    refs_start: int | None = None,
) -> list[Citation]:
    """Detect inline author-date citations like Author et al. (2023) or Author 2023."""
    citations = []

    for para_idx, para in enumerate(doc.paragraphs):
        if refs_start is not None and para_idx >= refs_start:
            break
        text = para.text
        if not text:
            continue

        # Pattern: Author (YYYY)
        for m in INLINE_AUTHOR_DATE_RE.finditer(text):
            if _overlaps_hyperlink(para_idx, m.start(), m.end(), hyperlink_spans):
                continue

            # Skip if this is part of a parenthetical we already caught
            # Check if the match is inside parentheses
            before = text[:m.start()].rstrip()
            if before.endswith("("):
                continue

            author = m.group(1).strip()
            year = m.group(2)

            if author in _FALSE_POSITIVE_AUTHORS:
                continue

            # Check for signal word before the citation
            signal = None
            signal_m = SIGNAL_RE.search(text[:m.start()])
            if signal_m:
                signal = signal_m.group(1)

            citations.append(Citation(
                id=next_id(),
                type="inline_author_date",
                display_text=m.group(0),
                paragraph_index=para_idx,
                surrounding_sentence=_get_surrounding_sentence(
                    text, m.start(), m.end()
                ),
                author_name=author,
                year=year,
                signal_word=signal,
            ))

        # Pattern: Author YYYY (no parens) — like "Hullman 2025 describes"
        for m in INLINE_NO_PARENS_RE.finditer(text):
            if _overlaps_hyperlink(para_idx, m.start(), m.end(), hyperlink_spans):
                continue

            # Avoid matching things already caught by INLINE_AUTHOR_DATE_RE
            author = m.group(1).strip()
            year = m.group(2)

            if author in _FALSE_POSITIVE_AUTHORS:
                continue

            # Skip if followed by closing paren (already caught as inline_author_date)
            after_pos = m.end()
            if after_pos < len(text) and text[after_pos] == ")":
                continue

            # Skip common false positives (years in non-citation context)
            # e.g. "in 2023" or "since 2020"
            before = text[:m.start()].rstrip()
            if before and before[-1] in "0123456789":
                continue

            citations.append(Citation(
                id=next_id(),
                type="inline_author_date",
                display_text=m.group(0),
                paragraph_index=para_idx,
                surrounding_sentence=_get_surrounding_sentence(
                    text, m.start(), m.end()
                ),
                author_name=author,
                year=year,
            ))

    return citations


def _detect_existing_footnotes(doc: Document, next_id) -> list[Citation]:
    """Detect existing footnotes that need Bluebook conversion."""
    citations = []

    for rel in doc.part.rels.values():
        if "footnotes" not in str(rel.reltype).lower():
            continue

        fn_part = rel.target_part
        fn_xml = etree.fromstring(fn_part.blob)

        for fn in fn_xml.findall(f".//{{{W}}}footnote"):
            fn_id_str = fn.get(f"{{{W}}}id")
            fn_type = fn.get(f"{{{W}}}type")

            # Skip separator footnotes
            if fn_type in ("separator", "continuationSeparator"):
                continue

            fn_id = int(fn_id_str) if fn_id_str else None

            # Get footnote text
            texts = [t.text for t in fn.findall(f".//{{{W}}}t") if t.text]
            fn_text = "".join(texts).strip()

            if not fn_text:
                continue

            # Extract URLs from footnote
            urls = URL_RE.findall(fn_text)
            url = urls[0] if urls else None

            # Check for hyperlinks in footnote
            for hl in fn.findall(f".//{{{W}}}hyperlink"):
                rid = hl.get(f"{{{R}}}id")
                if rid:
                    try:
                        hl_url = fn_part.rels[rid].target_ref
                        if not url:
                            url = hl_url
                    except KeyError:
                        pass

            citations.append(Citation(
                id=next_id(),
                type="existing_footnote",
                display_text=fn_text[:200],
                paragraph_index=-1,  # footnotes don't have a body paragraph
                surrounding_sentence=fn_text,
                url=url,
                existing_footnote_id=fn_id,
            ))

    return citations
