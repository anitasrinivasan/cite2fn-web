"""References section parsing and citation matching.

Parses the References/Bibliography/Works Cited section of a .docx document,
builds a lookup table, and matches inline citations to reference entries.
"""

from __future__ import annotations

import re
from docx import Document
from lxml import etree

from cite2fn.models import Citation, Reference

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def find_references_section(doc: Document) -> tuple[int | None, int | None]:
    """Find the start and end paragraph indices of the References section.

    Returns (start_index, end_index) where end_index is exclusive.
    Returns (None, None) if no references section is found.
    """
    start = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else ""
        is_heading = "heading" in style.lower() or "Heading" in style

        if text.lower() in ("references", "bibliography", "works cited"):
            if is_heading or len(text) < 20:
                start = i
                break

    if start is None:
        return None, None

    # End is the last paragraph or the next heading of equal/higher level
    end = len(doc.paragraphs)
    return start, end


def parse_references(doc_path: str) -> list[Reference]:
    """Parse all entries from the References section.

    Each reference entry is typically one paragraph. Some entries span
    multiple paragraphs (e.g., when title wraps).
    """
    doc = Document(doc_path)
    start, end = find_references_section(doc)
    if start is None:
        return []

    refs: list[Reference] = []
    current_text = ""
    current_para_idx = None

    for i in range(start + 1, end):
        text = doc.paragraphs[i].text.strip()
        if not text:
            # Empty paragraph might separate entries
            if current_text:
                ref = _parse_single_reference(current_text, current_para_idx)
                if ref:
                    refs.append(ref)
                current_text = ""
                current_para_idx = None
            continue

        # Heuristic: a new reference entry starts with an author name
        # (uppercase letter) and the previous entry has been accumulating text.
        if current_text and _looks_like_new_entry(text, current_text):
            ref = _parse_single_reference(current_text, current_para_idx)
            if ref:
                refs.append(ref)
            current_text = text
            current_para_idx = i
        else:
            if current_text:
                current_text += " " + text
            else:
                current_text = text
                current_para_idx = i

    # Don't forget the last entry
    if current_text:
        ref = _parse_single_reference(current_text, current_para_idx)
        if ref:
            refs.append(ref)

    # Resolve bookmark anchors to reference entries
    _resolve_anchors(doc, refs, start, end)

    return refs


def _looks_like_new_entry(text: str, prev_text: str) -> bool:
    """Heuristic: does this text look like the start of a new reference entry?"""
    # New entries typically start with an author last name
    # and contain a year in parentheses or after a colon
    if re.match(r"^[A-ZÀ-Ž]", text):
        # Check if prev_text looks complete (has a year and substantial content)
        if re.search(r"\d{4}", prev_text) and len(prev_text) > 30:
            return True
    return False


# Pattern: Author, A. B., Author2, C. (YYYY) or Author, A. B. (YYYY):
_REF_AUTHOR_YEAR = re.compile(
    r"^(.+?)\s*[\(\:]?\s*(\d{4})\s*[\)\:]?"
)

# Pattern for extracting first author's last name
_FIRST_AUTHOR = re.compile(r"^([A-ZÀ-Ž][a-zà-ž'\-]*(?:[A-Z][a-zà-ž'\-]+)*)")


def _parse_single_reference(text: str, para_idx: int | None) -> Reference | None:
    """Parse a single reference entry text into a Reference object."""
    if not text or len(text) < 10:
        return None

    # Extract year
    year_match = re.search(r"\((\d{4})\)", text)
    if not year_match:
        year_match = re.search(r"(\d{4})", text)
    year = year_match.group(1) if year_match else None

    # Extract authors (everything before the year or the first colon/period after names)
    authors = []
    first_author_match = _FIRST_AUTHOR.match(text)
    if first_author_match:
        authors.append(first_author_match.group(1))

    # Try to extract additional author last names from the author block
    # The author block is typically everything before the year
    if year_match:
        author_block = text[: year_match.start()]
    else:
        author_block = text[:60]  # First 60 chars as fallback

    # Find all capitalized surnames in the author block
    # Pattern: "LastName, F." or "LastName, F. F." or "and LastName"
    for m in re.finditer(r"(?:and\s+|&\s+)?([A-ZÀ-Ž][a-zà-ž'\-]+(?:[A-Z][a-zà-ž'\-]+)*)", author_block):
        name = m.group(1)
        if name not in authors and name not in ("The", "In", "A", "An"):
            authors.append(name)

    # Extract title (usually in quotes or after the year)
    title = None
    title_match = re.search(r'["\u201c](.+?)["\u201d]', text)
    if title_match:
        title = title_match.group(1)
    elif year_match:
        # Title is often after "YYYY): " or "YYYY: "
        after_year = text[year_match.end():]
        after_year = re.sub(r"^[\)\:\s]+", "", after_year)
        # Take up to the first period or comma that's followed by a journal/publisher
        title_end = re.search(r'[,\.]\s*(?:[A-Z]|http)', after_year)
        if title_end:
            title = after_year[: title_end.start()].strip().strip('""\u201c\u201d')

    return Reference(
        full_text=text,
        paragraph_index=para_idx or 0,
        authors=authors,
        year=year,
        title=title,
    )


def _resolve_anchors(
    doc: Document, refs: list[Reference], refs_start: int, refs_end: int
) -> None:
    """Find bookmark anchors within the References section and attach them to Reference objects.

    Google Docs exports use bookmarks like _hqe1036ucurq that point to specific
    paragraphs in the References section.
    """
    body = doc.element.body
    nsmap = {"w": W}

    # Build a map: paragraph_index -> list of bookmark names
    para_bookmarks: dict[int, list[str]] = {}
    for i, para in enumerate(doc.paragraphs):
        if i < refs_start or i >= refs_end:
            continue
        para_elem = para._element
        for bm_start in para_elem.findall(f".//{{{W}}}bookmarkStart"):
            name = bm_start.get(f"{{{W}}}name")
            if name and not name.startswith("_GoBack"):
                para_bookmarks.setdefault(i, []).append(name)

    # Match bookmarks to references by paragraph index
    for ref in refs:
        if ref.paragraph_index in para_bookmarks:
            ref.anchors = para_bookmarks[ref.paragraph_index]


def match_citations_to_references(
    citations: list[Citation], references: list[Reference]
) -> list[Citation]:
    """Match citations to reference entries.

    Updates citations in-place with matched_reference field.
    Returns the updated list.
    """
    if not references:
        return citations

    # Build lookup tables
    anchor_to_ref: dict[str, Reference] = {}
    for ref in references:
        for anchor in ref.anchors:
            anchor_to_ref[anchor] = ref

    # Fuzzy lookup by (author_last_name, year)
    author_year_to_ref: dict[tuple[str, str], list[Reference]] = {}
    for ref in references:
        if ref.year:
            for author in ref.authors:
                key = (author.lower(), ref.year)
                author_year_to_ref.setdefault(key, []).append(ref)

    for cite in citations:
        # Try anchor match first (most reliable)
        if cite.internal_anchor and cite.internal_anchor in anchor_to_ref:
            ref = anchor_to_ref[cite.internal_anchor]
            cite.matched_reference = ref.full_text
            continue

        # Try author+year match
        if cite.author_name and cite.year:
            # Normalize: take first word of author name
            author_key = cite.author_name.split()[0].lower()
            key = (author_key, cite.year)
            if key in author_year_to_ref:
                matches = author_year_to_ref[key]
                cite.matched_reference = matches[0].full_text
                continue

    return citations
