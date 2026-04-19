"""Document assembly pipeline.

Takes a document and formatted citations, produces the final output with:
- Footnotes, endnotes, or a list of references
- Bluebook or APA citation style
- Inline text cleaned (footnote/endnote modes only)
- Word comments for flagged issues (footnote/endnote modes only)
- Supra/id. short forms applied (Bluebook footnote/endnote only)
- References section removed (footnote/endnote modes only)
"""

from __future__ import annotations

import copy
import json
import re
from pathlib import Path
from docx import Document

from cite2fn.models import Citation, citations_from_json
from cite2fn.footnotes import (
    FootnoteManager,
    merge_adjacent_footnote_refs,
    normalize_footnote_placement,
)
from cite2fn.cleanup import classify_cleanup_rule, apply_cleanup
from cite2fn.comments import (
    add_no_source_comment,
    add_fetch_failed_comment,
    add_low_confidence_comment,
)
from cite2fn.supra import normalize_source_key, apply_short_forms
from cite2fn.docx_io import save_document, remove_references_section
from cite2fn.references_list import insert_references_list, sort_references

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def assemble_document(
    input_path: str,
    output_path: str,
    citations_data: list[dict],
    use_endnotes: bool = False,
    keep_references: bool = False,
    output_format: str = "footnotes",
    style: str = "bluebook",
) -> dict:
    """Assemble the final document with all conversions applied.

    Args:
        input_path: Path to the original .docx
        output_path: Path for the output .docx
        citations_data: List of citation dicts with bluebook_text filled in
        use_endnotes: Use endnotes instead of footnotes (legacy flag)
        keep_references: Don't remove the References section
        output_format: "footnotes", "endnotes", or "references"
        style: "bluebook" or "apa"

    Returns:
        Report dict with stats and issues.
    """
    # Normalize legacy flag
    if use_endnotes and output_format == "footnotes":
        output_format = "endnotes"

    doc = Document(input_path)
    citations = [Citation.from_dict(d) for d in citations_data]

    report = {
        "input": input_path,
        "output": output_path,
        "style": style,
        "output_format": output_format,
        "total_citations": len(citations),
        "footnotes_inserted": 0,
        "existing_footnotes_converted": 0,
        "references_listed": 0,
        "comments_added": 0,
        "references_removed": False,
        "issues": [],
    }

    # --- List of References mode ---
    if output_format == "references":
        return _assemble_references_list(doc, citations, output_path, report)

    # --- Footnotes / Endnotes mode ---
    use_endnotes_flag = output_format == "endnotes"
    fn_manager = FootnoteManager(doc, use_endnotes=use_endnotes_flag)

    # Separate existing footnotes from new citations
    existing_fn_cites = [c for c in citations if c.type == "existing_footnote"]
    new_cites = [c for c in citations if c.type != "existing_footnote"]

    # Sort new citations by paragraph index (process in document order)
    new_cites.sort(key=lambda c: (c.paragraph_index, c.id))

    # --- Step 1: Convert existing footnotes ---
    for cite in existing_fn_cites:
        if cite.bluebook_text and cite.existing_footnote_id is not None:
            fn_manager.replace_footnote_content(cite.existing_footnote_id, cite.bluebook_text)
            report["existing_footnotes_converted"] += 1

    # --- Step 2: Prepare footnotes for new citations ---
    # Build the supra/id. data structure
    footnote_entries = []

    for cite in new_cites:
        if not cite.bluebook_text:
            continue

        source_key = normalize_source_key(
            author=cite.author_name,
            url=cite.url,
        )

        footnote_entries.append({
            "citation_id": cite.id,
            "note_id": None,  # Will be assigned during insertion
            "bluebook_text": cite.bluebook_text,
            "source_key": source_key,
            "author_name": cite.author_name,
            "signal_word": cite.signal_word,
            "short_form_type": None,
        })

    # Apply supra/id. (Bluebook only — APA doesn't use short forms)
    temp_note_id = fn_manager._next_id
    for entry in footnote_entries:
        entry["note_id"] = temp_note_id
        temp_note_id += 1

    if style == "bluebook":
        footnote_entries = apply_short_forms(footnote_entries)

    # Build lookup from citation_id to formatted text
    cite_to_text: dict[str, str] = {}
    for entry in footnote_entries:
        cite_to_text[entry["citation_id"]] = entry["bluebook_text"]

    # --- Step 2b: Validate formatting markers ---
    if style == "bluebook":
        for entry in footnote_entries:
            warnings = _validate_bluebook_markers(entry["bluebook_text"])
            for w in warnings:
                report["issues"].append(f"Citation {entry['citation_id']}: {w}")

    # --- Step 3: Insert footnotes and clean text ---
    # Two-pass within each paragraph: insert-all-then-cleanup-all. Interleaving
    # cleanups with insertions (the old behavior) broke position lookup because
    # earlier cleanups deleted the text a later citation's search relied on.
    paragraphs_cites: dict[int, list[Citation]] = {}
    for cite in new_cites:
        paragraphs_cites.setdefault(cite.paragraph_index, []).append(cite)

    for para_idx in sorted(paragraphs_cites.keys(), reverse=True):
        para_cites = paragraphs_cites[para_idx]
        para = doc.paragraphs[para_idx]
        para_elem = para._element
        para_text = para.text

        def _first_occurrence(c: Citation) -> int:
            if c.display_text:
                idx = para_text.find(c.display_text)
                if idx >= 0:
                    return idx
            if c.year:
                idx = para_text.find(c.year)
                if idx >= 0:
                    return idx
            return 10**9

        para_cites.sort(key=_first_occurrence)

        # Capture run + hyperlink lists once and hold them through the loop so
        # lxml wrappers stay alive — their id() is only stable while they're
        # referenced. Recomputing inside _find_insert_position would invalidate
        # every claimed id between calls.
        top_runs = [r for r in para_elem if r.tag == f"{{{W}}}r"]
        all_runs = para_elem.findall(f".//{{{W}}}r")
        hyperlinks = para_elem.findall(f".//{{{W}}}hyperlink")

        claimed_ids: set[int] = set()
        classified: list[tuple[Citation, int]] = []
        for cite in reversed(para_cites):
            if cite.id not in cite_to_text:
                continue
            bluebook_text = cite_to_text[cite.id]

            rule = classify_cleanup_rule(cite, para_text)
            cite.cleanup_rule = rule
            classified.append((cite, rule))

            insert_after = _find_insert_position(
                cite, claimed_ids, top_runs, all_runs, hyperlinks
            )
            if insert_after is None:
                report["issues"].append(
                    f"Citation {cite.id}: could not anchor footnote — appended at paragraph tail"
                )
            fn_manager.insert_footnote(bluebook_text, para_elem, insert_after)
            report["footnotes_inserted"] += 1

        for cite, rule in classified:
            apply_cleanup(doc, cite, rule)

    # --- Step 3b: Merge adjacent footnote refs ---
    # Per Bluebook convention, multiple authorities cited at the same point
    # get a single footnote with sources joined by "; ". This also masks
    # position-lookup collisions where our code placed multiple refs next to
    # each other when they should have been spread across the paragraph.
    merged_total = 0
    for para in doc.paragraphs:
        merged_total += merge_adjacent_footnote_refs(para._element, fn_manager)
    if merged_total:
        report["footnotes_merged"] = merged_total

    # --- Step 3c: Normalize footnote placement across all paragraphs ---
    # Post-processing pass that fixes three defects insertion alone can't
    # guarantee: stray whitespace before a ref (from cleanup removing a
    # parenthetical year), FN before sentence-level punctuation, and FN
    # immediately followed by a letter without a space. Idempotent — no-op
    # on paragraphs already in correct shape.
    for para in doc.paragraphs:
        normalize_footnote_placement(para._element)

    # --- Step 4: Add comments for issues ---
    for cite in citations:
        if cite.type == "existing_footnote":
            continue

        para_idx = cite.paragraph_index
        if para_idx < 0 or para_idx >= len(doc.paragraphs):
            continue

        para = doc.paragraphs[para_idx]

        # No source URL
        if (cite.type in ("parenthetical", "inline_author_date")
                and not cite.url and not cite.matched_reference):
            add_no_source_comment(doc, para)
            report["comments_added"] += 1

        # Fetch failed
        if cite.fetched_metadata and cite.fetched_metadata.get("fetch_error"):
            add_fetch_failed_comment(doc, para, cite.url or "unknown")
            report["comments_added"] += 1

        # Low confidence
        if cite.confidence == "needs_review":
            add_low_confidence_comment(doc, para)
            report["comments_added"] += 1

    # --- Step 5: Remove references section ---
    if not keep_references:
        removed = remove_references_section(doc)
        if removed > 0:
            report["references_removed"] = True

    # --- Step 6: Save ---
    save_document(doc, output_path)

    return report


def _assemble_references_list(
    doc: Document,
    citations: list[Citation],
    output_path: str,
    report: dict,
) -> dict:
    """Assemble a document with a List of References at the end.

    Body text is left untouched. An alphabetically sorted reference list
    is appended at the end of the document.
    """
    # Remove old References section first (will be replaced)
    removed = remove_references_section(doc)
    if removed > 0:
        report["references_removed"] = True

    # Build entries for sorting
    entries = []
    for cite in citations:
        if not cite.bluebook_text:
            continue
        entries.append({
            "bluebook_text": cite.bluebook_text,
            "author_name": cite.author_name,
        })

    # Sort alphabetically and deduplicate
    sorted_refs = sort_references(entries)
    # Deduplicate (same formatted text)
    seen = set()
    unique_refs = []
    for ref in sorted_refs:
        if ref not in seen:
            seen.add(ref)
            unique_refs.append(ref)

    # Insert the reference list
    count = insert_references_list(doc, unique_refs)
    report["references_listed"] = count

    save_document(doc, output_path)
    return report


def _validate_bluebook_markers(bluebook_text: str) -> list[str]:
    """Flag potential formatting issues in bluebook_text."""
    warnings = []
    for match in re.finditer(r"~([^~]+)~", bluebook_text):
        inner = match.group(1).strip()
        # Strip punctuation and spaces to get the core text
        core = inner.replace("&", "").replace(" ", "").replace(".", "")
        # If the core is all-lowercase and short, it's likely an acronym
        # that shouldn't be in small caps at all
        if core.islower() and len(core) <= 5:
            warnings.append(
                f"Possible acronym in small caps: ~{inner}~ — "
                f"acronyms like {inner.upper()} should be plain uppercase, not in ~tildes~"
            )
    return warnings


def _split_text_element_after(run, t_elem, search_text):
    """Split `run` so that its text up to and including `search_text` stays in
    `run` (which becomes the "prefix"), and everything after `search_text`
    moves into a new "suffix" run inserted immediately after `run` in the
    paragraph.

    Returns the new suffix run, or None if the split is a no-op (search_text
    is already at the very end of the run's text).

    Assumes `search_text` appears in `t_elem.text`. The run is expected to have
    a single meaningful `w:t` (typical for body runs); for multi-`w:t` runs
    only the first split point is handled.
    """
    text = t_elem.text or ""
    idx = text.find(search_text)
    if idx < 0:
        return None
    split_pos = idx + len(search_text)
    if split_pos >= len(text):
        # Already at end — no suffix to carve out
        return None

    prefix_text = text[:split_pos]
    suffix_text = text[split_pos:]
    t_elem.text = prefix_text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    suffix_run = copy.deepcopy(run)
    # Strip all <w:t> content from the copy, then put the suffix into the first one
    suffix_t_elems = suffix_run.findall(f"{{{W}}}t")
    if not suffix_t_elems:
        return None
    for st in suffix_t_elems:
        st.text = ""
    suffix_t_elems[0].text = suffix_text
    suffix_t_elems[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    parent = run.getparent()
    parent.insert(list(parent).index(run) + 1, suffix_run)
    return suffix_run


def _find_insert_position(
    cite: Citation,
    claimed_ids: set[int],
    top_runs: list,
    all_runs: list,
    hyperlinks: list,
):
    """Find the XML element after which to insert the footnote reference.

    For hyperlinked citations: insert after the matching hyperlink element.
    For text citations: insert after a run containing the search text, preferring
    runs at the paragraph top level (runs inside <w:hyperlink> can be removed
    by a later cleanup for a different citation, which would drag our ref with
    them). `claimed_ids` tracks `id()` of elements already chosen for earlier
    citations in the same paragraph so two FNs don't collide on one anchor.

    The element lists must be captured ONCE in the caller and stay alive across
    all calls for a given paragraph — lxml reuses Python wrappers only while
    they're referenced, so recomputing inside this function would invalidate
    every id we've previously stored.

    Returns None if nothing can be anchored — caller appends to the tail.
    """
    if cite.type in ("hyperlink_external", "hyperlink_internal"):
        for hl in hyperlinks:
            if id(hl) in claimed_ids:
                continue
            texts = [t.text for t in hl.findall(f".//{{{W}}}t") if t.text]
            hl_text = "".join(texts)
            if cite.display_text.strip("()[].,; ") in hl_text or hl_text in cite.display_text:
                claimed_ids.add(id(hl))
                return hl

    # Search order: display_text first (most specific — lets us split cleanly
    # at the end of the full citation), then year, then author_name. Always
    # split on match so two citations whose anchor text lies within the same
    # big run still get distinct positions. claimed_ids is not used here —
    # the split guarantees distinct anchors even when two citations search
    # overlapping text.
    for search_text in (cite.display_text, cite.year, cite.author_name):
        if not search_text:
            continue
        for run in reversed(top_runs):
            for t in run.findall(f"{{{W}}}t"):
                if not (t.text and search_text in t.text):
                    continue
                suffix = _split_text_element_after(run, t, search_text)
                if suffix is not None:
                    top_runs.append(suffix)
                    all_runs.append(suffix)
                claimed_ids.add(id(run))
                return run
        # Fallback: text lives inside a hyperlink. Return the CONTAINING
        # hyperlink as the anchor so the FN is inserted as a paragraph-level
        # sibling of the hyperlink — not as a child of it. This guarantees
        # the FN survives even if a *different* citation's rule 2/3 cleanup
        # removes that hyperlink. We deliberately do NOT add the hyperlink to
        # claimed_ids: multiple citations may legitimately anchor there, and
        # the adjacent-ref merger will consolidate them afterwards.
        for run in reversed(all_runs):
            if run in top_runs:
                continue
            for t in run.findall(f"{{{W}}}t"):
                if t.text and search_text in t.text:
                    anc = run.getparent()
                    while anc is not None and anc.tag != f"{{{W}}}hyperlink":
                        if anc.tag == f"{{{W}}}p":
                            anc = None
                            break
                        anc = anc.getparent()
                    if anc is not None and anc.tag == f"{{{W}}}hyperlink":
                        return anc
                    break  # non-hyperlink ancestor — skip this text

    return None
